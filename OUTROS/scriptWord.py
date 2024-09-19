import os
from docx import Document
import openpyxl


# Função para extrair título e primeiro parágrafo de arquivos Word
def extrair_titulo_paragrafo(arquivo_word):
    try:
        doc = Document(arquivo_word)
        titulo = doc.paragraphs[0].text if len(doc.paragraphs) > 0 else ""
        primeiro_paragrafo = doc.paragraphs[1].text if len(doc.paragraphs) > 1 else ""
        return titulo, primeiro_paragrafo
    except Exception as e:
        print(f"Erro ao processar o arquivo {arquivo_word}: {e}")
        return "", ""


# Função para processar os arquivos e extrair o título e o primeiro parágrafo
def processar_arquivos_em_diretorio(diretorio):
    arquivos_processados = []

    # Itera sobre todos os arquivos no diretório
    for arquivo in os.listdir(diretorio):
        if arquivo.endswith('.docx'):
            caminho_arquivo = os.path.join(diretorio, arquivo)

            # Verifica se o arquivo realmente existe
            if os.path.exists(caminho_arquivo):
                titulo, paragrafo = extrair_titulo_paragrafo(caminho_arquivo)
                arquivos_processados.append((arquivo, titulo, paragrafo))
            else:
                print(f"Arquivo não encontrado: {caminho_arquivo}")

    return arquivos_processados


# Função para salvar os resultados em um arquivo Excel
def salvar_resultados_no_excel(arquivos_processados, arquivo_excel):
    # Criando ou abrindo o arquivo Excel
    try:
        workbook = openpyxl.load_workbook(arquivo_excel)
    except FileNotFoundError:
        workbook = openpyxl.Workbook()
        workbook.active.append(['Arquivo', 'Título', 'Primeiro Parágrafo'])

    sheet = workbook.active

    # Salvando os dados dos arquivos processados
    for arquivo, titulo, paragrafo in arquivos_processados:
        sheet.append([arquivo, titulo, paragrafo])

    # Salvando o arquivo Excel
    workbook.save(arquivo_excel)
    print(f"Resultados salvos no arquivo {arquivo_excel}")


# Exemplo de uso
diretorio = r'C:\Users\erics\OneDrive\FECAP\2024\09. Setembro'
arquivos_processados = processar_arquivos_em_diretorio(diretorio)

# Salvar os resultados em um arquivo Excel
arquivo_excel = 'resultados_arquivos_word.xlsx'
salvar_resultados_no_excel(arquivos_processados, arquivo_excel)

# Exibindo os resultados (opcional)
for arquivo, titulo, paragrafo in arquivos_processados:
    print(f"Arquivo: {arquivo}\nTítulo: {titulo}\nParágrafo: {paragrafo}\n")
